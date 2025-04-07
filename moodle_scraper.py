import streamlit as st
import os
import time
import io # Added for in-memory file handling
import google.generativeai as genai
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from docx import Document # Added for Word document creation
from docx.shared import Inches # Added for potential formatting
from dotenv import load_dotenv # Added



# --- Load Environment Variables ---
load_dotenv() # Load variables from .env file

# --- Helper Function for Logging ---
def log_status(message, level="info"):
    """Helper function to log status messages with consistent formatting."""
    if level == "success":
        st.success(message)
    elif level == "warning":
        st.warning(message)
    elif level == "error":
        st.error(message)
    else:  # default to info
        st.info(message)
    # Optionally print to console during development
    # print(f"[{level.upper()}] {message}")

# --- Configuration ---
MOODLE_LOGIN_URL = os.getenv("MOODLE_LOGIN_URL", "https://moodle.spit.ac.in/login/index.php") # Keep default if not set

# Load credentials from environment variables
USERNAME = os.getenv("MOODLE_USERNAME")
PASSWORD = os.getenv("MOODLE_PASSWORD")

# Load user details from environment variables
YOUR_NAME = os.getenv("USER_NAME")
YOUR_UID = os.getenv("USER_UID")

# Load Gemini key - can still be overridden by secrets if preferred
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyDqUAINl5B7Qy52s2_ukiIl3noELeRMHFc") # Keep hardcoded as fallback

# --- Validate that required variables are loaded ---
if not USERNAME or not PASSWORD:
    st.error("Moodle username or password not found in environment variables (.env file). Please set MOODLE_USERNAME and MOODLE_PASSWORD.")
    st.stop()
if not YOUR_NAME or not YOUR_UID:
    st.error("User name or UID not found in environment variables (.env file). Please set USER_NAME and USER_UID.")
    st.stop()
if not GEMINI_API_KEY:
    st.error("Gemini API Key not configured. Please set GEMINI_API_KEY in .env or script.")
    st.stop()

# Optional: Specify path if WebDriver is not in system PATH
# WEBDRIVER_PATH = os.getenv("WEBDRIVER_PATH")

# --- Locators ---
USERNAME_FIELD_LOCATOR = (By.ID, "username")
PASSWORD_FIELD_LOCATOR = (By.ID, "password")
LOGIN_BUTTON_LOCATOR = (By.ID, "loginbtn")
# This locator might need adjustment based on the specific Moodle page structure.
ASSIGNMENT_DESCRIPTION_LOCATOR = (By.CSS_SELECTOR, ".no-overflow") # <-- VERIFY AND UPDATE THIS

# --- Function to Scrape and Analyze ---
@st.cache_data(show_spinner=False) # Cache results based on URL for faster re-runs
def get_assignment_details(_assignment_url): # Removed name and uid parameters
    """Logs into Moodle, navigates to assignment, scrapes description, and sends to Gemini with custom prompt."""
    # Access variables loaded from environment
    name = YOUR_NAME
    uid = YOUR_UID
    driver = None
    description_text = None
    gemini_response_text = None
    error_message = None
    status_updates = []

    try:
        st.info("Setting up WebDriver...")
        options = webdriver.ChromeOptions()
        # options.add_argument("--headless") # Optional: Run in background
        options.add_argument("--start-maximized")
        # When running in Streamlit Cloud, you might need additional options:
        # options.add_argument('--disable-gpu')
        # options.add_argument('--no-sandbox')
        driver = webdriver.Chrome(options=options)
        wait = WebDriverWait(driver, 20) # Increased wait time slightly

        # --- Login ---
        st.info(f"Navigating to login page: {MOODLE_LOGIN_URL}")
        driver.get(MOODLE_LOGIN_URL)

        st.info("Locating login elements...")
        username_field = wait.until(EC.presence_of_element_located(USERNAME_FIELD_LOCATOR))
        password_field = wait.until(EC.presence_of_element_located(PASSWORD_FIELD_LOCATOR))
        login_button = wait.until(EC.element_to_be_clickable(LOGIN_BUTTON_LOCATOR))

        st.info("Entering credentials...")
        username_field.send_keys(USERNAME)
        password_field.send_keys(PASSWORD)

        st.info("Attempting to log in...")
        login_button.click()

        # --- Wait for Login Confirmation ---
        try:
            wait.until(EC.presence_of_element_located((By.ID, "page-site-index"))) # Adjust if needed
            st.success("Login successful.")
        except TimeoutException:
            error_message = "Login failed or took too long. Check credentials, network, or if the confirmation element ID changed."
            try:
                error_element = driver.find_element(By.ID, "loginerrormessage")
                error_message += f" Moodle Error: {error_element.text}"
            except NoSuchElementException:
                pass
            # driver.save_screenshot("login_failure.png") # Debugging
            raise Exception(error_message)

        # --- Navigate to Assignment ---
        st.info(f"Navigating to assignment page: {_assignment_url}")
        driver.get(_assignment_url)

        # --- Extract Description ---
        st.info(f"Locating assignment description using locator: {ASSIGNMENT_DESCRIPTION_LOCATOR}...")
        time.sleep(2) # Small delay

        try:
            description_element = wait.until(EC.visibility_of_element_located(ASSIGNMENT_DESCRIPTION_LOCATOR)) # Wait for visibility too
            description_text = description_element.get_attribute('innerText').strip()
            if not description_text:
                 log_status("Description element found, but 'innerText' was empty. Trying '.text' attribute...", "warning")
                 description_text = description_element.text.strip()

            if description_text:
                log_status("Assignment description extracted.", "success")

                # --- Close Browser Immediately After Extraction ---
                if driver:
                    log_status("Closing browser after description extraction.")
                    driver.quit()
                    driver = None # Set driver to None so finally block doesn't try again

                # --- Send Description to Gemini with Custom Prompt ---
                log_status("Sending description to Gemini with custom prompt...")
                try:
                    genai.configure(api_key=GEMINI_API_KEY)
                    model = genai.GenerativeModel('gemini-2.0-flash')

                    # --- Construct the new prompt using hardcoded name/uid --- 
                    prompt = f"""Solve the given below experiment and give me solution along with aim ,thoery ,conclusion and all in total of 1500 words and also include my name {name} uid {uid}
The problem statement is:

{description_text}
"""
                    log_status(f"Generated Prompt for Gemini (first 200 chars): \n{prompt[:200]}...", "info")

                    response = model.generate_content(prompt)

                    try:
                         gemini_response_text = response.text
                    except ValueError: # Correctly indented except block
                        # Handle cases where the response might be blocked
                        gemini_response_text = f"Gemini response blocked or unavailable. Finish Reason: {response.prompt_feedback.block_reason}. Parts: {response.parts}"
                        st.warning(gemini_response_text)

                    st.success("Received response from Gemini.")
                except Exception as gemini_e:
                    error_message = f"Error communicating with Gemini: {gemini_e}"
                    st.error(error_message)
                    # Keep description_text even if Gemini fails
            else:
                error_message = f"Could not extract text from the description element located by {ASSIGNMENT_DESCRIPTION_LOCATOR}. The element might be empty or the locator incorrect."
                st.warning(error_message)
                # driver.save_screenshot("description_not_found.png") # Debugging

        except TimeoutException:
            error_message = f"Error: Could not find the assignment description element using locator {ASSIGNMENT_DESCRIPTION_LOCATOR} within the time limit. Verify the locator."
            st.error(error_message)
            # driver.save_screenshot("description_locator_fail.png") # Debugging
        except Exception as desc_e:
            error_message = f"An unexpected error occurred while extracting the description: {desc_e}"
            st.error(error_message)
            # driver.save_screenshot("description_error.png") # Debugging

    except TimeoutException as e:
        error_message = f"A timeout occurred: {e}. Check element locators, network speed, or increase wait time."
        st.error(error_message)
        # if driver: driver.save_screenshot("timeout_error.png") # Debugging
    except NoSuchElementException as e:
         error_message = f"Element not found: {e}. One of the locators might be incorrect."
         st.error(error_message)
         # if driver: driver.save_screenshot("element_not_found_error.png") # Debugging
    except Exception as e:
        error_message = f"An unexpected error occurred: {e}"
        st.error(error_message)
        # if driver: driver.save_screenshot("unexpected_error.png") # Debugging
    finally:
        if driver: # Check if driver is still active before quitting
            log_status("Closing browser in finally block.")
            driver.quit()

    return description_text, gemini_response_text, error_message, status_updates

# --- Function to Create Word Document ---
def create_word_document(description, analysis):
    """Creates a Word document containing the description and analysis."""
    document = Document()
    document.add_heading('Moodle Assignment Analysis', 0)

    document.add_heading('Assignment Description', level=1)
    if description:
        document.add_paragraph(description)
    else:
        document.add_paragraph("[No description extracted]")

    document.add_paragraph() # Add a space

    document.add_heading('Gemini Analysis', level=1)
    if analysis:
        # Basic handling: add paragraph by paragraph. Could be enhanced for lists/formatting.
        # For now, just adding the whole text. Markdown won't render directly.
        document.add_paragraph(analysis)
    else:
        document.add_paragraph("[No analysis available]")

    # Save document to a BytesIO object
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0) # Rewind the stream to the beginning
    return file_stream

# --- Streamlit App ---
st.set_page_config(layout="wide", page_title="Moodle Assignment Analyzer")
st.title("📝 Moodle Assignment Scraper & Analyzer")
st.caption("Enter a Moodle assignment URL to extract its description and get an analysis from Gemini.")

st.sidebar.header("🔒 Configuration Status")
# Display status based on loaded env vars
st.sidebar.text_input("Moodle Username", value=USERNAME if USERNAME else "Not Set", key="moodle_user_disp", disabled=True)
st.sidebar.text_input("Moodle Password", value="********" if PASSWORD else "Not Set", key="moodle_pass_disp", type="password", disabled=True)
st.sidebar.text_input("User Name", value=YOUR_NAME if YOUR_NAME else "Not Set", key="user_name_disp", disabled=True)
st.sidebar.text_input("User UID", value=YOUR_UID if YOUR_UID else "Not Set", key="user_uid_disp", disabled=True)
st.sidebar.text_input("Gemini API Key", value="********" if GEMINI_API_KEY else "Not Set", key="gemini_key_disp", type="password", disabled=True)

st.divider()

st.header("Enter Assignment URL")
assignment_url_input = st.text_input(
    "Paste the full URL of the Moodle assignment page:",
    # Example URL placeholder
    placeholder="e.g., https://moodle.yourinstitution.com/mod/assign/view.php?id=12345"
)

if st.button("Scrape and Analyze Assignment", key="analyze_button", type="primary"):
    if assignment_url_input and assignment_url_input.startswith(("http://", "https://")):
        st.session_state.results = None # Clear previous results
        st.session_state.status = []   # Clear previous status
        status_placeholder = st.empty()
        with st.spinner("🚀 Processing... Logging in, navigating, scraping, and analyzing. This may take up to a minute."):
            # --- Pass name and UID to the function ---
            description, gemini_response, error, status_log = get_assignment_details(assignment_url_input)
            st.session_state.results = (description, gemini_response, error)
            st.session_state.status = status_log

        if error and not (description or gemini_response): # If total failure
             st.error(f"Failed to process the assignment. Error: {error}")
        else: # Partial or full success
            st.header("Results")
            col1, col2 = st.columns(2)

            with col1:
                st.subheader("Extracted Assignment Description")
                if description:
                    st.text_area("Description", description, height=400)
                elif error:
                     st.warning(f"Could not extract description. Error details: {error}")
                else:
                    st.warning("No description was extracted (element might be empty or locator incorrect).")

            with col2:
                st.subheader("🧠 Gemini Analysis")
                if gemini_response:
                    st.markdown(gemini_response, unsafe_allow_html=True) # Allow basic HTML if needed
                elif error and description and not gemini_response: # Description extracted but Gemini failed
                    st.warning(f"Could not get Gemini analysis. Details in log.", icon="⚠️")
                elif not description:
                     st.info("Cannot analyze description as it was not extracted.", icon="ℹ️")
                else: # No error, description exists, but no Gemini response (e.g., API issue)
                     st.info("No Gemini analysis available. Check processing log.", icon="ℹ️")

            # --- Download Button ---
            if description or gemini_response: # Only show button if there's something to download
                st.divider()
                st.subheader("⬇️ Download Results")
                try:
                    doc_stream = create_word_document(description, gemini_response)
                    st.download_button(
                        label="Download as Word (.docx)",
                        data=doc_stream,
                        file_name="moodle_assignment_analysis.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as doc_e:
                    st.error(f"Error creating Word document: {doc_e}")

            # Display non-fatal error prominently if it occurred alongside some results
            if error and (description or gemini_response):
                 st.warning(f"Processing completed with issues: {error}", icon="⚠️")

    elif not assignment_url_input:
        st.warning("Please enter a Moodle assignment URL.")
    else:
        st.error("Invalid URL format. Please enter a valid HTTP/HTTPS URL.")

# Hide the original script execution logic if run directly
# The Streamlit commands above handle the execution flow.
# The original try/except/finally block outside a function is removed. 